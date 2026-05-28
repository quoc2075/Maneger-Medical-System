"""Xuất báo cáo tài chính — PDF, Word, Excel, CSV."""
from __future__ import annotations

import csv
import io
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from django.conf import settings
from django.http import HttpResponse


THEME = {
    'primary': '0077B6',   # xanh y tế
    'accent': '00A8CC',
    'light': 'F8FAFC',
    'muted': 'E2E8F0',
    'text': '0F172A',
    'danger': 'C1121F',
}

DEFAULT_BRAND = {
    'clinic_name': 'PHÒNG KHÁM ĐA KHOA',
    'hotline': 'Hotline: 1900 0000',
    'email': 'Email: lienhe@phongkham.vn',
    'website': 'Website: https://phongkham.vn',
    'address': 'Địa chỉ: 123 Đường ABC, TP.HCM',
    'logo_path': '',
}


def _get_brand() -> Dict[str, str]:
    brand = dict(DEFAULT_BRAND)
    cfg = getattr(settings, 'BAOCAO_BRAND', None)
    if isinstance(cfg, dict):
        for k in brand:
            v = cfg.get(k)
            if v:
                brand[k] = str(v).strip()
    return brand


def _logo_candidates(brand_logo_path: str) -> List[Path]:
    base = Path(settings.BASE_DIR)
    out = []
    if brand_logo_path:
        p = Path(brand_logo_path)
        if not p.is_absolute():
            p = base / brand_logo_path
        out.append(p)
    out += [
        base / 'baocao' / 'assets' / 'logo.png',
        base / 'baocao' / 'assets' / 'logo.jpg',
        base / 'baocao' / 'assets' / 'logo.jpeg',
    ]
    return out


def _pick_logo_path(brand: Dict[str, str]) -> Optional[str]:
    for p in _logo_candidates(brand.get('logo_path', '')):
        if p.is_file():
            return str(p)
    return None


def _pdf_font_name() -> str:
    """Đăng ký font Unicode (tiếng Việt) nếu có trên máy chủ."""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        if 'VNFont' in pdfmetrics.getRegisteredFontNames():
            return 'VNFont'
        candidates = [
            os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'arial.ttf'),
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        ]
        for path in candidates:
            if os.path.isfile(path):
                pdfmetrics.registerFont(TTFont('VNFont', path))
                return 'VNFont'
    except Exception:
        pass
    return 'Helvetica'


def _verify_url(data: dict) -> str:
    tu = data.get('tu') or ''
    den = data.get('den') or ''
    nhom = data.get('nhom') or 'ngay'
    return f"https://verify.phongkham.vn/financial?tu={tu}&den={den}&nhom={nhom}"


def _meta(data: dict, request) -> Dict[str, str]:
    now = datetime.now().strftime('%d/%m/%Y %H:%M')
    brand = _get_brand()
    return {
        'co_quan_chu_quan': 'BỘ Y TẾ',
        'don_vi': brand['clinic_name'],
        'ten_bao_cao': 'BÁO CÁO THỐNG KÊ TÀI CHÍNH VÀ DOANH THU',
        'ky_bao_cao': data.get('ky_bao_cao') or f"{data.get('tu')} — {data.get('den')}",
        'nguoi_xuat': getattr(request.user, 'ho_ten', '') or getattr(request.user, 'ten_dang_nhap', ''),
        'thoi_diem_xuat': now,
        'so_hieu_bieu_mau': 'BM-BV-TC-01',
        'nhom': data.get('nhom') or 'ngay',
        'bo_loc': f"tu={data.get('tu') or ''}; den={data.get('den') or ''}; nhom={data.get('nhom') or ''}",
        'hotline': brand['hotline'],
        'email': brand['email'],
        'website': brand['website'],
        'address': brand['address'],
        'logo_path': _pick_logo_path(brand) or '',
        'verify_url': _verify_url(data),
    }


def _money(v) -> str:
    try:
        return f'{float(v or 0):,.0f}'
    except (TypeError, ValueError):
        return str(v)


def _money_num(v) -> float:
    try:
        return float(v or 0)
    except (TypeError, ValueError):
        return 0.0


def _summary_rows(data: dict) -> List[Tuple[str, str]]:
    return [
        ('Tổng doanh thu', _money(data.get('doanh_thu'))),
        ('Doanh thu — đơn hàng (app/quầy)', _money(data.get('doanh_thu_don_hang'))),
        ('Doanh thu — đơn thuốc toa', _money(data.get('doanh_thu_don_thuoc'))),
        ('Doanh thu — tiêm chủng', _money(data.get('doanh_thu_tiem'))),
        ('Giá vốn ước tính', _money(data.get('gia_von_uoc_tinh'))),
        ('Lợi nhuận ước tính', _money(data.get('loi_nhuan_uoc_tinh'))),
        ('Số giao dịch', str(data.get('so_giao_dich') or 0)),
    ]


def _detail_headers(nhom: str) -> List[str]:
    if nhom == 'bac_si':
        return [
            'Mã BS',
            'Bác sĩ',
            'Đơn hàng',
            'Đơn toa',
            'Tiêm',
            'Tổng DT',
            'SL đơn',
            'SL toa',
            'SL tiêm',
        ]
    if nhom == 'thang':
        return ['Tháng', 'Đơn hàng', 'Đơn toa', 'Tiêm', 'Tổng']
    if nhom == 'nam':
        return ['Năm', 'Đơn hàng', 'Đơn toa', 'Tiêm', 'Tổng']
    return ['Ngày', 'Đơn hàng', 'Đơn toa', 'Tiêm', 'Tổng']


def _detail_rows(data: dict) -> List[List[str]]:
    nhom = data.get('nhom') or 'ngay'
    rows = data.get('bang_chi_tiet') or data.get('theo_ngay') or []
    out: List[List[str]] = []
    if nhom == 'bac_si':
        for r in rows:
            out.append([
                r.get('ma_bac_si') or '',
                r.get('ten_bac_si') or '',
                _money(r.get('doanh_thu_don_hang')),
                _money(r.get('doanh_thu_don_thuoc')),
                _money(r.get('doanh_thu_tiem')),
                _money(r.get('tong_doanh_thu')),
                str(r.get('so_don_hang') or 0),
                str(r.get('so_don_thuoc') or 0),
                str(r.get('so_lan_tiem') or 0),
            ])
        return out
    label_key = {'thang': 'thang', 'nam': 'nam'}.get(nhom, 'ngay')
    for r in rows:
        out.append([
            r.get(label_key) or r.get('ngay') or '',
            _money(r.get('doanh_thu_don_hang')),
            _money(r.get('doanh_thu_don_thuoc')),
            _money(r.get('doanh_thu_tiem')),
            _money(r.get('tong_doanh_thu')),
        ])
    return out


def _detail_rows_numeric(data: dict) -> List[List]:
    nhom = data.get('nhom') or 'ngay'
    rows = data.get('bang_chi_tiet') or data.get('theo_ngay') or []
    out = []
    if nhom == 'bac_si':
        for r in rows:
            out.append([
                r.get('ma_bac_si') or '',
                r.get('ten_bac_si') or '',
                _money_num(r.get('doanh_thu_don_hang')),
                _money_num(r.get('doanh_thu_don_thuoc')),
                _money_num(r.get('doanh_thu_tiem')),
                _money_num(r.get('tong_doanh_thu')),
                int(r.get('so_don_hang') or 0),
                int(r.get('so_don_thuoc') or 0),
                int(r.get('so_lan_tiem') or 0),
            ])
        return out
    label_key = {'thang': 'thang', 'nam': 'nam'}.get(nhom, 'ngay')
    for r in rows:
        out.append([
            r.get(label_key) or r.get('ngay') or '',
            _money_num(r.get('doanh_thu_don_hang')),
            _money_num(r.get('doanh_thu_don_thuoc')),
            _money_num(r.get('doanh_thu_tiem')),
            _money_num(r.get('tong_doanh_thu')),
        ])
    return out


def _filename(data: dict, ext: str) -> str:
    tu = (data.get('tu') or 'tu').replace('-', '')
    den = (data.get('den') or 'den').replace('-', '')
    nhom = data.get('nhom') or 'ngay'
    return f'bao-cao-tai-chinh_{tu}_{den}_{nhom}.{ext}'


def export_financial_csv(data: dict, request) -> HttpResponse:
    meta = _meta(data, request)
    buf = io.StringIO(newline='')
    w = csv.writer(buf, quoting=csv.QUOTE_MINIMAL, lineterminator='\n')
    w.writerow([meta['co_quan_chu_quan']])
    w.writerow([meta['don_vi']])
    w.writerow([meta['ten_bao_cao']])
    w.writerow(['Biểu mẫu', meta['so_hieu_bieu_mau']])
    w.writerow(['Kỳ báo cáo', meta['ky_bao_cao']])
    w.writerow(['Nhóm chi tiết', meta['nhom']])
    w.writerow(['Bộ lọc', meta['bo_loc']])
    w.writerow(['Người xuất', meta['nguoi_xuat']])
    w.writerow(['Thời điểm xuất', meta['thoi_diem_xuat']])
    w.writerow([])
    w.writerow(['Chỉ tiêu', 'Giá trị'])
    for label, val in _summary_rows(data):
        w.writerow([label, val])
    w.writerow([])
    w.writerow(_detail_headers(meta['nhom']))
    for row in _detail_rows(data):
        w.writerow(row)
    resp = HttpResponse('\ufeff' + buf.getvalue(), content_type='text/csv; charset=utf-8')
    resp['Content-Disposition'] = f'attachment; filename="{_filename(data, "csv")}"'
    return resp


def export_financial_xlsx(data: dict, request) -> HttpResponse:
    try:
        from openpyxl import Workbook
        from openpyxl.chart import BarChart, Reference
        from openpyxl.formatting.rule import ColorScaleRule
        from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
        from openpyxl.utils import get_column_letter
    except ImportError:
        return HttpResponse(
            '{"error":"Thiếu thư viện openpyxl. Chạy: pip install openpyxl"}',
            content_type='application/json',
            status=503,
        )

    meta = _meta(data, request)
    wb = Workbook()
    ws = wb.active
    ws.title = 'Tong hop'

    thin = Side(style='thin', color=THEME['muted'])
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    fill_header = PatternFill('solid', fgColor=THEME['primary'])
    fill_alt = PatternFill('solid', fgColor=THEME['light'])
    fill_total = PatternFill('solid', fgColor='E0F2FE')
    white_bold = Font(color='FFFFFF', bold=True)

    ws.merge_cells('A1:F1')
    ws['A1'] = meta['co_quan_chu_quan']
    ws['A1'].font = Font(bold=True, size=12, color=THEME['primary'])
    ws['A1'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:F2')
    ws['A2'] = meta['don_vi']
    ws['A2'].font = Font(bold=True, size=14, color=THEME['text'])
    ws['A2'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:F3')
    ws['A3'] = meta['ten_bao_cao']
    ws['A3'].font = Font(bold=True, size=15, color=THEME['accent'])
    ws['A3'].alignment = Alignment(horizontal='center')

    ws.append([])
    ws.append(['Biểu mẫu', meta['so_hieu_bieu_mau'], 'Kỳ báo cáo', meta['ky_bao_cao'], 'Nhóm', meta['nhom']])
    ws.append(['Người xuất', meta['nguoi_xuat'], 'Ngày xuất', meta['thoi_diem_xuat'], 'Bộ lọc', meta['bo_loc']])

    for r in (5, 6):
        for c in range(1, 7):
            ws.cell(row=r, column=c).border = border
            ws.cell(row=r, column=c).alignment = Alignment(vertical='center', wrap_text=True)

    ws.append([])
    ws.append(['Chỉ tiêu', 'Giá trị'])
    sum_header_row = ws.max_row
    ws[f'A{sum_header_row}'].font = white_bold
    ws[f'B{sum_header_row}'].font = white_bold
    ws[f'A{sum_header_row}'].fill = fill_header
    ws[f'B{sum_header_row}'].fill = fill_header
    ws[f'A{sum_header_row}'].alignment = Alignment(horizontal='center')
    ws[f'B{sum_header_row}'].alignment = Alignment(horizontal='center')

    for label, val in _summary_rows(data):
        ws.append([label, _money_num(val.replace(',', '')) if val.replace(',', '').isdigit() else val])
    sum_end_row = ws.max_row

    for r in range(sum_header_row + 1, sum_end_row + 1):
        ws[f'A{r}'].border = border
        ws[f'B{r}'].border = border
        ws[f'B{r}'].number_format = '#,##0'
        if 'Tổng doanh thu' in str(ws[f'A{r}'].value):
            ws[f'A{r}'].font = Font(bold=True, color=THEME['primary'])
            ws[f'B{r}'].font = Font(bold=True, color=THEME['primary'])
            ws[f'A{r}'].fill = fill_total
            ws[f'B{r}'].fill = fill_total

    ws.append([])
    headers = _detail_headers(meta['nhom'])
    ws.append(headers)
    detail_header_row = ws.max_row
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=detail_header_row, column=c)
        cell.font = white_bold
        cell.fill = fill_header
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    detail_rows = _detail_rows_numeric(data)
    for row in detail_rows:
        ws.append(row)
    detail_start = detail_header_row + 1
    detail_end = ws.max_row

    for r in range(detail_start, detail_end + 1):
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=r, column=c)
            cell.border = border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
            if c > 1 and isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0'
        if (r - detail_start) % 2 == 1:
            for c in range(1, len(headers) + 1):
                ws.cell(row=r, column=c).fill = fill_alt

    total_col = len(headers)
    if detail_end >= detail_start:
        ws.auto_filter.ref = f"A{detail_header_row}:{get_column_letter(total_col)}{detail_end}"
        ws.conditional_formatting.add(
            f"{get_column_letter(total_col)}{detail_start}:{get_column_letter(total_col)}{detail_end}",
            ColorScaleRule(
                start_type='min',
                start_color='E2E8F0',
                mid_type='percentile',
                mid_value=50,
                mid_color='BAE6FD',
                end_type='max',
                end_color='0284C7',
            ),
        )

    ws.freeze_panes = f"A{detail_header_row + 1}"
    for col in range(1, len(headers) + 1):
        max_len = 12
        for r in range(1, ws.max_row + 1):
            v = ws.cell(row=r, column=col).value
            max_len = max(max_len, len(str(v or '')))
        ws.column_dimensions[get_column_letter(col)].width = min(max_len + 2, 42)

    ws2 = wb.create_sheet('Top thuoc')
    ws2.append(['Nguồn', 'Tên thuốc', 'SL', 'Doanh thu'])
    for x in data.get('top_thuoc_don_hang') or []:
        ws2.append(['Đơn hàng', x.get('ten_thuoc'), int(x.get('so_luong') or 0), _money_num(x.get('doanh_thu'))])
    for x in data.get('top_thuoc_theo_toa') or []:
        ws2.append(['Theo toa', x.get('ten_thuoc'), int(x.get('so_luong') or 0), _money_num(x.get('doanh_thu'))])
    for c in range(1, 5):
        ws2.cell(row=1, column=c).font = white_bold
        ws2.cell(row=1, column=c).fill = fill_header
    ws2.freeze_panes = 'A2'
    ws2.auto_filter.ref = f"A1:D{ws2.max_row}"

    ws_chart = wb.create_sheet('Bieu do')
    ws_chart.append(['Mục', 'Doanh thu'])
    for label, value in _summary_rows(data)[:4]:
        ws_chart.append([label, _money_num(value.replace(',', '')) if value.replace(',', '').isdigit() else 0])

    chart = BarChart()
    chart.title = 'Tổng hợp doanh thu'
    chart.y_axis.title = 'VNĐ'
    chart.x_axis.title = 'Chỉ tiêu'
    vals = Reference(ws_chart, min_col=2, min_row=1, max_row=5)
    cats = Reference(ws_chart, min_col=1, min_row=2, max_row=5)
    chart.add_data(vals, titles_from_data=True)
    chart.set_categories(cats)
    chart.width = 18
    chart.height = 8
    ws_chart.add_chart(chart, 'D2')

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    resp = HttpResponse(
        bio.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    resp['Content-Disposition'] = f'attachment; filename="{_filename(data, "xlsx")}"'
    return resp


def export_financial_docx(data: dict, request) -> HttpResponse:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        return HttpResponse(
            '{"error":"Thiếu thư viện python-docx. Chạy: pip install python-docx"}',
            content_type='application/json',
            status=503,
        )

    meta = _meta(data, request)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.columns[0].width = Inches(1.2)
    header_tbl.columns[1].width = Inches(5.8)
    if meta['logo_path']:
        try:
            header_tbl.cell(0, 0).paragraphs[0].add_run().add_picture(meta['logo_path'], width=Inches(0.9))
        except Exception:
            header_tbl.cell(0, 0).text = ''
    info_p = header_tbl.cell(0, 1).paragraphs[0]
    info_p.add_run(meta['don_vi']).bold = True
    info_p.add_run(f"\n{meta['hotline']} | {meta['email']}")
    info_p.add_run(f"\n{meta['website']}")
    info_p.add_run(f"\n{meta['address']}")

    title = doc.add_paragraph(meta['ten_bao_cao'])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(16)

    doc.add_paragraph(f"Biểu mẫu: {meta['so_hieu_bieu_mau']}")
    doc.add_paragraph(f"Kỳ báo cáo: {meta['ky_bao_cao']}")
    doc.add_paragraph(f"Bộ lọc dữ liệu: {meta['bo_loc']}")
    doc.add_paragraph(f"Người xuất: {meta['nguoi_xuat']} | Ngày xuất: {meta['thoi_diem_xuat']}")

    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Table Grid'
    t1.rows[0].cells[0].text = 'Chỉ tiêu'
    t1.rows[0].cells[1].text = 'Giá trị'
    for label, val in _summary_rows(data):
        row = t1.add_row().cells
        row[0].text = label
        row[1].text = val

    doc.add_paragraph()
    doc.add_heading('Bảng chi tiết', level=2)
    hdr = _detail_headers(meta['nhom'])
    t2 = doc.add_table(rows=1, cols=len(hdr))
    t2.style = 'Table Grid'
    for i, h in enumerate(hdr):
        t2.rows[0].cells[i].text = h
    for row in _detail_rows(data):
        cells = t2.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.cell(0, 0).text = 'Người lập báo cáo\n(ký, ghi rõ họ tên)'
    sig.cell(0, 1).text = 'Đại diện phòng khám\n(ký, đóng dấu)'

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    resp = HttpResponse(
        bio.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    resp['Content-Disposition'] = f'attachment; filename="{_filename(data, "docx")}"'
    return resp


def export_financial_pdf(data: dict, request) -> HttpResponse:
    try:
        from reportlab.graphics.barcode import createBarcodeDrawing
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        return HttpResponse(
            '{"error":"Thiếu thư viện reportlab. Chạy: pip install reportlab"}',
            content_type='application/json',
            status=503,
        )

    try:
        meta = _meta(data, request)
        font = _pdf_font_name()
        bio = io.BytesIO()
        doc = SimpleDocTemplate(
            bio,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=100,
            bottomMargin=70,
        )
        styles = getSampleStyleSheet()
        normal = ParagraphStyle('VN', parent=styles['Normal'], fontName=font, fontSize=10, textColor=colors.HexColor(f"#{THEME['text']}"))
        title = ParagraphStyle('TitleVN', parent=styles['Heading1'], fontName=font, fontSize=16, alignment=1, textColor=colors.HexColor(f"#{THEME['primary']}"))
        subtitle = ParagraphStyle('SubVN', parent=styles['Normal'], fontName=font, fontSize=9, textColor=colors.HexColor('#334155'))

        def _decorate_page(canvas, _doc):
            canvas.saveState()
            w, h = A4
            logo = meta.get('logo_path')

            # Header cố định.
            if logo:
                try:
                    canvas.drawImage(logo, 38, h - 78, width=44, height=44, preserveAspectRatio=True, mask='auto')
                except Exception:
                    pass
            canvas.setFont(font, 11)
            canvas.setFillColor(colors.HexColor(f"#{THEME['primary']}"))
            canvas.drawString(88, h - 50, meta['don_vi'])
            canvas.setFont(font, 9)
            canvas.setFillColor(colors.HexColor('#475569'))
            canvas.drawString(88, h - 64, f"{meta['hotline']} | {meta['email']}")
            canvas.drawString(88, h - 76, f"{meta['website']} | {meta['address']}")
            canvas.setStrokeColor(colors.HexColor(f"#{THEME['muted']}"))
            canvas.line(36, h - 84, w - 36, h - 84)

            # Watermark logo mờ.
            if logo:
                try:
                    if hasattr(canvas, 'setFillAlpha'):
                        canvas.setFillAlpha(0.08)
                    canvas.drawImage(logo, w / 2 - 120, h / 2 - 120, width=240, height=240, preserveAspectRatio=True, mask='auto')
                    if hasattr(canvas, 'setFillAlpha'):
                        canvas.setFillAlpha(1)
                except Exception:
                    pass

            # Footer có số trang.
            canvas.setFont(font, 8)
            canvas.setFillColor(colors.HexColor('#64748B'))
            canvas.drawCentredString(w / 2, 26, f"Trang {_doc.page}")
            canvas.drawRightString(w - 36, 26, f"Xuất lúc: {meta['thoi_diem_xuat']}")
            canvas.restoreState()

        story = [
            Paragraph(meta['ten_bao_cao'], title),
            Spacer(1, 6),
            Paragraph(
                f"Kỳ báo cáo: {meta['ky_bao_cao']}<br/>"
                f"Bộ lọc: {meta['bo_loc']}<br/>"
                f"Người xuất: {meta['nguoi_xuat']} | Nhóm: {meta['nhom']}",
                subtitle,
            ),
            Spacer(1, 14),
        ]

        sum_data = [['Chỉ tiêu', 'Giá trị']] + [[a, b] for a, b in _summary_rows(data)]
        t1 = Table(sum_data, colWidths=[270, 220])
        t1.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(f"#{THEME['primary']}")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{THEME['muted']}")),
            ('FONTNAME', (0, 0), (-1, -1), font),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(f"#{THEME['light']}")),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor(f"#{THEME['text']}")),
            ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
            ('BACKGROUND', (0, 1), (1, 1), colors.HexColor('E0F2FE')),
            ('FONTNAME', (0, 1), (1, 1), font),
        ]))
        story.append(t1)
        story.append(Spacer(1, 14))

        hdr = _detail_headers(meta['nhom'])
        detail = [hdr] + _detail_rows(data)
        col_width = 540 / max(len(hdr), 5)
        t2 = Table(detail, repeatRows=1, colWidths=[col_width] * len(hdr))
        style_cmds = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(f"#{THEME['accent']}")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor(f"#{THEME['muted']}")),
            ('FONTNAME', (0, 0), (-1, -1), font),
            ('FONTSIZE', (0, 0), (-1, -1), 8.4),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('WORDWRAP', (0, 0), (-1, -1), 'CJK'),
            ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ]
        for i in range(1, len(detail)):
            if i % 2 == 0:
                style_cmds.append(('BACKGROUND', (0, i), (-1, i), colors.HexColor(f"#{THEME['light']}")))
        if len(hdr) >= 5:
            style_cmds.append(('TEXTCOLOR', (len(hdr) - 1, 1), (len(hdr) - 1, -1), colors.HexColor(f"#{THEME['primary']}")))
            style_cmds.append(('FONTNAME', (len(hdr) - 1, 1), (len(hdr) - 1, -1), font))
        t2.setStyle(TableStyle(style_cmds))
        story.append(t2)
        story.append(Spacer(1, 14))

        # Tổng kết + QR + chữ ký + dấu mộc giả lập.
        qr = createBarcodeDrawing('QR', value=meta['verify_url'], width=72, height=72, barLevel='M')
        final_tbl = Table(
            [
                [
                    Paragraph(
                        f"<b>Tổng kết cuối trang</b><br/>"
                        f"Tổng doanh thu: <b>{_money(data.get('doanh_thu'))} VNĐ</b><br/>"
                        f"Lợi nhuận ước tính: <b>{_money(data.get('loi_nhuan_uoc_tinh'))} VNĐ</b>",
                        normal,
                    ),
                    qr,
                    Paragraph(
                        "NGƯỜI LẬP BÁO CÁO<br/><br/>"
                        "(ký, ghi rõ họ tên)<br/><br/><br/>"
                        f"{meta['nguoi_xuat'] or '........................'}"
                        "<br/><font color='#C1121F'>[ĐÃ KÝ - DẤU ĐIỆN TỬ]</font>",
                        ParagraphStyle('Sign', parent=normal, alignment=1),
                    ),
                ]
            ],
            colWidths=[270, 90, 180],
        )
        final_tbl.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor(f"#{THEME['muted']}")),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (1, 0), (1, 0), 'CENTER'),
            ('ALIGN', (2, 0), (2, 0), 'CENTER'),
            ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('F8FBFF')),
        ]))
        story.append(final_tbl)

        doc.build(story, onFirstPage=_decorate_page, onLaterPages=_decorate_page)
        bio.seek(0)
        resp = HttpResponse(bio.read(), content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{_filename(data, "pdf")}"'
        return resp
    except Exception as exc:
        return HttpResponse(
            f'{{"error":"Lỗi tạo PDF: {str(exc)[:200]}"}}',
            content_type='application/json',
            status=500,
        )


EXPORTERS = {
    'csv': export_financial_csv,
    'xlsx': export_financial_xlsx,
    'excel': export_financial_xlsx,
    'docx': export_financial_docx,
    'word': export_financial_docx,
    'pdf': export_financial_pdf,
}


def export_financial(data: dict, request, fmt: str) -> HttpResponse:
    key = (fmt or '').lower().strip()
    fn = EXPORTERS.get(key)
    if not fn:
        return HttpResponse(
            '{"error":"dinh_dang phải là pdf, docx, xlsx hoặc csv"}',
            content_type='application/json',
            status=400,
        )
    try:
        return fn(data, request)
    except Exception as exc:
        return HttpResponse(
            f'{{"error":"Lỗi xuất file: {str(exc)[:200]}"}}',
            content_type='application/json',
            status=500,
        )
